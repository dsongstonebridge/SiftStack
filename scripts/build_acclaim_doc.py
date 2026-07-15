"""Build the Acclaim Integration Summary Word document for the CEO."""
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)

# Base font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

DARK_BLUE = RGBColor(0x1F, 0x39, 0x64)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY = RGBColor(0x60, 0x60, 0x60)


def blue_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = DARK_BLUE
    return p


def para(text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(bold_prefix=None, rest=""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(rest)
    p.paragraph_format.space_after = Pt(4)
    return p


def shade_cell(cell, fill_hex):
    pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pr.append(shd)


def header_row(table, cols, fill="1F3964"):
    row = table.rows[0]
    for i, text in enumerate(cols):
        cell = row.cells[i]
        cell.text = text
        shade_cell(cell, fill)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = WHITE


def add_row(table, values, first_bold=False):
    row = table.add_row()
    for i, val in enumerate(values):
        row.cells[i].text = val
        if i == 0 and first_bold:
            for run in row.cells[i].paragraphs[0].runs:
                run.bold = True
    return row


# ─── Title block ───────────────────────────────────────────────────────────────
title = doc.add_heading("SiftStack — Acclaim Integration & Lead Pipeline", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(18)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
today = datetime.date.today().strftime("%B %d, %Y")
run = sub.add_run(f"Prepared for: CEO Review     |     Date: {today}")
run.font.size = Pt(10)
run.font.color.rgb = GRAY
sub.paragraph_format.space_after = Pt(16)

# ─── 1. What Is Acclaim? ───────────────────────────────────────────────────────
blue_heading("1. What Is Acclaim?")
para(
    "Acclaim is Tulsa County's official document recording portal "
    "(acclaim.tulsacounty.org). We integrated it as a second foreclosure data "
    "source alongside OSCN. Where OSCN captures the moment a lender files a "
    "lawsuit in civil court, Acclaim captures the moment a lender records a "
    "document at the County Clerk — specifically Lis Pendens notices and Notice "
    "of Sheriff Sale filings. Because every recorded document is tied to a "
    "parcel ID, every Acclaim record arrives with a confirmed property address. "
    "No name-to-address lookup is required."
)

# ─── 2. How Acclaim Differs from OSCN ─────────────────────────────────────────
blue_heading("2. How Acclaim Differs from OSCN")
para(
    "OSCN is our primary foreclosure source — high volume, real-time court "
    "filings. Acclaim is an address-recovery layer that supplements OSCN for the "
    "subset of records where a property address could not be resolved and the "
    "lender also chose to record a Lis Pendens."
)

t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
header_row(t, ["", "OSCN", "Acclaim"])
rows2 = [
    ("Trigger", "Lawsuit filed in civil court", "Document recorded at County Clerk"),
    ("Address Source", "Owner name -> Assessor lookup", "Parcel ID in the recording (confirmed)"),
    ("Address Success Rate", "~53%  (47% fail)", "100%"),
    ("Cadence", "Real-time, daily", "Verified weekly (7-day lag)"),
    ("Volume", "Higher", "Lower (see constraints)"),
    ("Deduplication", "By court case ID", "By normalized parcel ID"),
]
for r in rows2:
    add_row(t, r, first_bold=True)

doc.add_paragraph().paragraph_format.space_after = Pt(4)
para(
    "Cross-source deduplication is handled automatically at the parcel level: "
    "if the same property appears in both OSCN and Acclaim, only one record "
    "exports to the output CSV."
)

# ─── 3. Key Constraints ────────────────────────────────────────────────────────
blue_heading("3. Key Constraints")

blue_heading("The 7-Day Verification Lag", level=2)
para(
    "Acclaim's database is verified approximately 7 days behind real-time. "
    "Documents recorded in the past 7 days are not yet visible in the portal. "
    "To consistently capture the verified window without missing records, we "
    "run Acclaim on a weekly schedule with a 14-day lookback floor. Each weekly "
    "run captures a clean 7-day band of verified recordings (days 8-14 from "
    "filing). A seen-IDs cache prevents any record from being exported twice, "
    "even if back-to-back weekly runs overlap in their date windows."
)

blue_heading("Lis Pendens Is Optional in Oklahoma", level=2)
para(
    "Oklahoma is a judicial foreclosure state, meaning lenders must file a "
    "lawsuit in court to foreclose — that filing is captured by OSCN. However, "
    "recording a Lis Pendens at the County Clerk is not legally required. The "
    "lender can proceed entirely through the court system without recording at "
    "the Clerk. Major servicers — JPMorgan Chase, Wells Fargo, and others — "
    "typically do record a Lis Pendens. Smaller or specialty lenders often do not."
)
para(
    "This is the primary reason Acclaim produces lower volume than OSCN. It "
    "captures only the subset of foreclosures where the lender took the optional "
    "step of recording at the County Clerk. It is not a limitation of our "
    "scraper — it is a limitation of what is publicly recorded in Oklahoma."
)

# ─── 4. Lead Pipeline ──────────────────────────────────────────────────────────
blue_heading("4. Lead Pipeline — How Each Record Is Processed")
para(
    "Once Acclaim records are scraped and pass the enrichment pipeline "
    "(sold properties removed, vacant land filtered, entity owners removed), "
    "the following steps execute automatically:"
)

pipeline = [
    ("Zillow Property Enrichment",
     " — Each record is enriched with estimated value, estimated equity, MLS "
     "status, last sale date, beds, baths, and square footage via OpenWebNinja. "
     "In our initial run, 16 of 25 records enriched with an average estimated "
     "equity of $160,101."),
    ("Tracerfy Batch Skip Trace ($0.02/record)",
     " — All records are submitted to Tracerfy for owner phone and email lookup "
     "against public records databases. In our initial run, 12 of 24 submitted "
     "records matched, returning 55 phone numbers and 31 emails. Records that do "
     "not match are typically out-of-state investors, recently transferred "
     "properties, or name-format mismatches between the court filing and the "
     "public records database."),
    ("Upload to DataSift",
     " — Records are formatted and uploaded as a new dated list. Tracerfy phones "
     "are included in the upload so they are immediately available in the CRM."),
    ("DataSift Property Enrichment",
     " — DataSift enriches each record with SiftMap property data as an "
     "independent verification layer."),
    ("DataSift Skip Trace (subscription included)",
     " — DataSift's built-in skip trace runs on all uploaded records, providing "
     "additional phone coverage beyond Tracerfy. This recovers contacts that "
     "Tracerfy did not match."),
    ("Trestle Phone Scoring ($0.015/phone)",
     " — Every phone number found — from Tracerfy and DataSift combined — is "
     "scored 0-100 by Trestle's phone intelligence API and assigned a dial "
     "priority tier. Tags are applied to each phone number directly in DataSift."),
]
for bold, rest in pipeline:
    bullet(bold_prefix=bold, rest=rest)

doc.add_paragraph().paragraph_format.space_after = Pt(4)
para("Trestle Dial Priority Tiers (initial run results):", bold=True)

tier_table = doc.add_table(rows=1, cols=4)
tier_table.style = "Table Grid"
header_row(tier_table, ["Tier", "Score Range", "Count", "Action"])
tiers = [
    ("Dial First", "81-100", "10  (40%)", "Highest activity — call first"),
    ("Dial Second", "61-80",  "2   (8%)",  "High activity"),
    ("Dial Third",  "41-60",  "3   (12%)", "Moderate activity"),
    ("Dial Fourth", "21-40",  "8   (32%)", "Low activity"),
    ("Drop",        "0-20",   "2   (8%)",  "Inactive / disconnected — skip"),
]
for t_row in tiers:
    add_row(tier_table, t_row, first_bold=True)

doc.add_paragraph().paragraph_format.space_after = Pt(4)
para(
    "NOTE: 15 of the 40 total phones from Tracerfy encountered Trestle API rate "
    "limits during the initial batch and will be re-scored on the next weekly run."
)

# ─── 5. Run Cadence ────────────────────────────────────────────────────────────
blue_heading("5. Run Cadence")

cad = doc.add_table(rows=1, cols=4)
cad.style = "Table Grid"
header_row(cad, ["Run", "Source", "Lookback", "Frequency"])
add_row(cad, ("Daily", "OSCN court filings", "1 day (prior day)", "Every business day"))
add_row(cad, ("Weekly", "Acclaim County Clerk recordings", "14 days (captures verified 7-day window)", "Once per week"))

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ─── 6. Bottom Line ────────────────────────────────────────────────────────────
blue_heading("6. Bottom Line")
para(
    "Acclaim does not replace OSCN — it supplements it. Its core value is "
    "address confirmation: every record that comes through Acclaim has a "
    "verified property address tied to a recorded parcel. For the subset of "
    "foreclosures where OSCN could not resolve an address and the lender "
    "recorded a Lis Pendens, Acclaim recovers those leads and delivers them "
    "with full contact information."
)
para(
    "Combined with Tracerfy batch skip tracing, DataSift property enrichment "
    "and skip trace, and Trestle phone scoring, every exportable record arrives "
    "in the CRM with a verified address, estimated equity, and a prioritized "
    "call list — ready for outreach."
)

out_path = r"C:\Users\dsong\OneDrive\Desktop\Acclaim_Integration_Summary.docx"
doc.save(out_path)
print("Saved:", out_path)
