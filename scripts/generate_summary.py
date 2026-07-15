"""Generate Word document summary for June 22, 2026 run."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

title = doc.add_heading('SiftStack Daily Run Summary', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Date: June 22, 2026')
doc.add_paragraph('Prepared by: Jeff Hoemann / SiftStack Automation')
doc.add_paragraph('')

# ── Executive Summary ──
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'Completed a combined daily OSCN scrape and weekly Acclaim scrape for Tulsa County. '
    'After enrichment and deduplication, 88 unique leads were uploaded to DataSift with '
    'full property data, skip trace results, and dial-priority phone scoring. '
    'Several pipeline issues were discovered and fixed during the run.'
)

# ── Scrape Results ──
doc.add_heading('Scrape Results', level=1)

doc.add_heading('Weekly Acclaim Run (June 8-22, 14-day lookback)', level=2)
t = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
for i, (k, v) in enumerate([
    ('NOTICE (Lis Pendens)', '35 records'),
    ('DECREE (Foreclosure Decrees)', '1 record'),
    ('Total scraped', '36 records'),
    ('Addresses found (Tulsa Assessor)', '31/36 (86%)'),
    ('Entity-owned removed', '2 (LLC/Inc)'),
    ('Final Acclaim leads', '29 foreclosure'),
]):
    t.cell(i, 0).text = k
    t.cell(i, 1).text = v

doc.add_paragraph('')
doc.add_heading('Daily OSCN Run (June 12-22)', level=2)
t = doc.add_table(rows=7, cols=2, style='Light Grid Accent 1')
for i, (k, v) in enumerate([
    ('OSCN Foreclosure', '139 cases'),
    ('OSCN Probate', '10 cases'),
    ('Tulsa World Legal Notices', '1 notice'),
    ('Total raw', '150 records'),
    ('Addresses found (Tulsa Assessor)', '85/150 (57%)'),
    ('After dedup + entity filter', '69 leads'),
    ('Breakdown', '61 foreclosure + 8 probate'),
]):
    t.cell(i, 0).text = k
    t.cell(i, 1).text = v

doc.add_paragraph('')
doc.add_heading('Combined Results', level=2)
t = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
for i, (k, v) in enumerate([
    ('Total before dedup', '98 records (29 Acclaim + 69 OSCN)'),
    ('Cross-source duplicates removed', '2 (same address, different source)'),
    ('Sold properties removed', '8 (Zillow flagged as recently sold)'),
    ('Final unique leads', '88 (80 foreclosure + 8 probate)'),
    ('Address overlap', '2 properties appeared in both Acclaim and OSCN'),
]):
    t.cell(i, 0).text = k
    t.cell(i, 1).text = v

# ── Enrichment ──
doc.add_heading('Enrichment Results', level=1)

doc.add_heading('Smarty Address Standardization', level=2)
doc.add_paragraph(
    'Smarty ran on all 96 records but returned 0 confirmed matches. '
    'Root cause: a hardcoded Tennessee state filter was rejecting all Oklahoma addresses. '
    'The log message "Smarty returned OK" was the state abbreviation, not a status. '
    'Fixed during session - will work correctly on the next run.'
)

doc.add_heading('Zillow Property Enrichment (OpenWebNinja)', level=2)
t = doc.add_table(rows=4, cols=2, style='Light Grid Accent 1')
for i, (k, v) in enumerate([
    ('Records enriched', '82/96 (85%)'),
    ('API calls used', '87 of 100 monthly quota'),
    ('Average estimated equity', '$199,898'),
    ('Remaining monthly quota', '13 calls'),
]):
    t.cell(i, 0).text = k
    t.cell(i, 1).text = v

doc.add_heading('Tracerfy Skip Trace', level=2)
doc.add_paragraph(
    'Initial run: 6/96 matched (6%) - caused by a name format bug. Acclaim records stored names '
    'as "LAST FIRST" without commas, so first/last were swapped when sent to Tracerfy.'
)
doc.add_paragraph(
    'After fix: 50/82 matched (61%) with 172 phones and 90 emails.'
)
t = doc.add_table(rows=4, cols=2, style='Light Grid Accent 1')
for i, (k, v) in enumerate([
    ('First run (bad names)', '$1.92 - 6 matches, 26 phones, 15 emails'),
    ('Second run (fixed names)', '$1.64 - 50 matches, 172 phones, 90 emails'),
    ('Total Tracerfy cost', '$3.56 (should have been $1.64)'),
    ('Wasted due to bug', '$1.92'),
]):
    t.cell(i, 0).text = k
    t.cell(i, 1).text = v

doc.add_heading('Trestle Phone Scoring', level=2)
doc.add_paragraph(
    'Trestle was run multiple times due to pipeline ordering issues. '
    'The flow has been corrected so Trestle only runs once (after DataSift skip trace) on future runs.'
)
t = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
for i, (k, v) in enumerate([
    ('Pipeline run 1 (wasted - bad names)', '198 phones - ~$2.97'),
    ('Pipeline run 2 (fixed names)', '198 phones - ~$2.97'),
    ('phone-validate (after DataSift)', '490 phones - $7.35'),
    ('Total Trestle cost', '~$13.29'),
    ('Ideal cost (single run)', '$7.35'),
]):
    t.cell(i, 0).text = k
    t.cell(i, 1).text = v

doc.add_heading('Final Dial Priority Distribution', level=2)
t = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
for i, (k, v) in enumerate([
    ('Dial First (best)', '227 phones'),
    ('Dial Second', '38 phones'),
    ('Dial Third', '31 phones'),
    ('Dial Fourth', '151 phones'),
    ('Drop (do not call)', '43 phones'),
    ('Total scored', '490 phones'),
]):
    t.cell(i, 0).text = k
    t.cell(i, 1).text = v

# ── DataSift ──
doc.add_heading('DataSift Upload', level=1)
doc.add_paragraph(
    'Records uploaded to DataSift. DataSift enrichment and skip trace were triggered successfully. '
    'Phone dial-priority tags uploaded via phone-validate. '
    'Each record is tagged with its source (Acclaim or OSCN), notice type, and Trestle dial tier.'
)

# ── Issues ──
doc.add_heading('Issues Discovered & Fixed', level=1)

issues = [
    ('Acclaim Kendo Grid returning 0 records',
     'AJAX timing race condition. Fixed with widget init waits, value verification, and auto-retry.'),
    ('BOKF (Bank of Oklahoma) records dropped',
     '5 foreclosure filings silently dropped. "BOKF" added to lender regex.'),
    ('Tracerfy 6% match rate',
     'Acclaim names in "LAST FIRST" format without commas caused first/last swap. Fixed with comma conversion.'),
    ('Smarty rejecting all Oklahoma addresses',
     'Hardcoded state filter (!= "TN") rejected every OK address. Fixed to compare against notice state.'),
    ('Trestle scores not saved to CSV',
     'Dial-priority scores computed but never written to Tags column. Fixed.'),
    ('DataSift tag automation blocked by overlay',
     'Modal overlay div intercepted clicks. Fixed by adding overlay removal.'),
    ('DataSift select-all finding 0 records',
     'Records not loaded when enrichment tried to select. Fixed with 60-second retry loop.'),
    ('Trestle running before DataSift skip trace',
     'Pipeline scored phones before DataSift found additional numbers. Restructured to run Trestle only via phone-validate after DataSift.'),
]
for title_text, desc in issues:
    p = doc.add_paragraph()
    r = p.add_run(title_text + ': ')
    r.bold = True
    p.add_run(desc)

# ── Cost Summary ──
doc.add_heading('Cost Summary', level=1)
t = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
for i, h in enumerate(['Service', 'Actual Cost', 'Ideal Cost']):
    t.cell(0, i).text = h
    t.cell(0, i).paragraphs[0].runs[0].bold = True
for i, (svc, actual, ideal) in enumerate([
    ('Zillow (OpenWebNinja)', '87 API calls', '87 API calls'),
    ('Tracerfy Skip Trace', '$3.56', '$1.64'),
    ('Trestle Phone Scoring', '~$13.29', '$7.35'),
    ('DataSift Enrichment', 'Included in plan', 'Included in plan'),
    ('DataSift Skip Trace', 'Included in plan', 'Included in plan'),
    ('Total Extra Cost (bugs)', '~$7.86', '$0.00'),
], 1):
    t.cell(i, 0).text = svc
    t.cell(i, 1).text = actual
    t.cell(i, 2).text = ideal

# ── Deep Prospecting Note ──
doc.add_heading('Note: Deep Prospecting / Obituary Enrichment', level=1)
doc.add_paragraph(
    'The 8 probate records were not fully enriched with heir/decision-maker data. '
    'The obituary enrichment pipeline requires three additional APIs:'
)
for b in [
    'Serper (Google Search API) - finds obituary URLs',
    'Firecrawl (JS web scraper) - scrapes obituary page text',
    'Anthropic/Claude API - extracts heir names, relationships, and locations from obituary text',
]:
    doc.add_paragraph(b, style='List Bullet')
doc.add_paragraph(
    'Without these, probate records have phone numbers for the deceased person rather than '
    'the executor/heir who would handle the property. An Anthropic API key is needed to enable this.'
)

# ── Next Steps ──
doc.add_heading('Recommended Next Steps', level=1)
for i, step in enumerate([
    'Activate Smarty US Street Address API subscription - address standardization will work on next run',
    'Set up Anthropic API key for probate deep prospecting (obituary/heir enrichment)',
    'Run phone-validate after each DataSift skip trace to score new phones only',
    'Future runs will be smoother - all bugs from today are fixed',
], 1):
    doc.add_paragraph(f'{i}. {step}')

doc.save('C:/Users/dsong/OneDrive/Desktop/SiftStack_Run_Summary_2026-06-22.docx')
print('Saved summary document')
